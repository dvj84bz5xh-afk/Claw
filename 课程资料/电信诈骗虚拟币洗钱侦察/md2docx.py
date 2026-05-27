"""
Markdown to Word Converter
将教纲和教案Markdown文件转换为Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def set_chinese_font(run, font_name='宋体', size=12, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

def md_to_docx(md_file, output_file, title=None):
    """将Markdown转换为Word文档"""
    doc = Document()
    
    # 设置默认中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    table_data = []
    in_table = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 处理表格
        if line.startswith('|') and line.endswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_data.append(cells)
            continue
        else:
            if in_table and len(table_data) > 2:
                # 创建表格（跳过表头分隔行）
                rows = [r for r in table_data if not all(c.replace('-','').replace(':','')=='' for c in r)]
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    for i, row in enumerate(rows):
                        for j, cell_text in enumerate(row):
                            table.rows[i].cells[j].text = cell_text
                table_data = []
            in_table = False
        
        # 处理标题
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            set_chinese_font(run, '黑体', 18, True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            set_chinese_font(run, '黑体', 16, True)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            set_chinese_font(run, '黑体', 14, True)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(line[5:])
            set_chinese_font(run, '黑体', 12, True)
        # 处理列表
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            for run in p.runs:
                set_chinese_font(run, '宋体', 12)
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
            for run in p.runs:
                set_chinese_font(run, '宋体', 12)
        # 普通段落
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                set_chinese_font(run, '宋体', 12)
    
    # 添加页脚
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "刑侦总队六支队  电信诈骗案中虚拟币洗钱侦察课程"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(output_file)
    print(f"已生成: {output_file}")

# 转换教纲
md_to_docx(
    '《电信诈骗案中虚拟币洗钱侦察》教纲.md',
    '《电信诈骗案中虚拟币洗钱侦察》教纲.docx',
    '《电信诈骗案中虚拟币洗钱侦察》课程教纲'
)

# 转换教案
md_to_docx(
    '《电信诈骗案中虚拟币洗钱侦察》教案.md',
    '《电信诈骗案中虚拟币洗钱侦察》教案.docx',
    '《电信诈骗案中虚拟币洗钱侦察》课程教案'
)

print("所有Word文档已生成完成！")
